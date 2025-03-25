"""
Document Processor Module

This module handles the extraction of text from various document formats including:
- PDF files
- DOCX files
- CSV files
- Image files (for OCR)
- Video files (frame extraction and OCR)
"""

import os
import PyPDF2
import docx
import pandas as pd
import pytesseract
import cv2
from moviepy import VideoFileClip
import numpy as np
from PIL import Image
import logging

# Configure logging
logging.basicConfig(level=logging.INFO, format='%(asctime)s - %(levelname)s - %(message)s')
logger = logging.getLogger(__name__)

def process_pdf(file_path):
    """
    Extract text from a PDF file.
    
    Args:
        file_path (str): Path to the PDF file
        
    Returns:
        str: Extracted text content
    """
    logger.info(f"Processing PDF: {file_path}")
    text = ""
    try:
        with open(file_path, 'rb') as file:
            pdf_reader = PyPDF2.PdfReader(file)
            for page_num in range(len(pdf_reader.pages)):
                page_text = pdf_reader.pages[page_num].extract_text()
                if page_text:
                    text += page_text + "\n\n"
        logger.info(f"Successfully extracted text from PDF: {file_path}")
        return text
    except Exception as e:
        logger.error(f"Error processing PDF {file_path}: {str(e)}")
        raise

def process_docx(file_path):
    """
    Extract text from a DOCX file.
    
    Args:
        file_path (str): Path to the DOCX file
        
    Returns:
        str: Extracted text content
    """
    logger.info(f"Processing DOCX: {file_path}")
    try:
        doc = docx.Document(file_path)
        text = ""
        for paragraph in doc.paragraphs:
            if paragraph.text:
                text += paragraph.text + "\n"
        
        # Process tables if present
        for table in doc.tables:
            for row in table.rows:
                row_text = ""
                for cell in row.cells:
                    row_text += cell.text + " | "
                text += row_text.strip(" | ") + "\n"
        
        logger.info(f"Successfully extracted text from DOCX: {file_path}")
        return text
    except Exception as e:
        logger.error(f"Error processing DOCX {file_path}: {str(e)}")
        raise

def process_csv(file_path):
    """
    Extract text representation from a CSV file.
    
    Args:
        file_path (str): Path to the CSV file
        
    Returns:
        str: Text representation of the CSV data
    """
    logger.info(f"Processing CSV: {file_path}")
    try:
        df = pd.read_csv(file_path)
        
        # Get column information
        columns_info = "Columns: " + ", ".join(df.columns) + "\n\n"
        
        # Convert DataFrame to a readable text format
        # Include column names and row data
        rows_text = df.to_string(index=False)
        
        # Add some basic statistics
        stats = "\n\nSummary Statistics:\n"
        numeric_columns = df.select_dtypes(include=['number']).columns
        if not numeric_columns.empty:
            stats += df[numeric_columns].describe().to_string()
        
        text = columns_info + rows_text + stats
        logger.info(f"Successfully processed CSV: {file_path}")
        return text
    except Exception as e:
        logger.error(f"Error processing CSV {file_path}: {str(e)}")
        raise

def process_image(file_path):
    """
    Extract text from an image using OCR.
    
    Args:
        file_path (str): Path to the image file
        
    Returns:
        str: Extracted text from the image
    """
    logger.info(f"Processing Image: {file_path}")
    try:
        # Read the image
        image = cv2.imread(file_path)
        
        # Preprocess the image for better OCR results
        # Convert to grayscale
        gray = cv2.cvtColor(image, cv2.COLOR_BGR2GRAY)
        
        # Apply thresholding
        _, thresh = cv2.threshold(gray, 150, 255, cv2.THRESH_BINARY | cv2.THRESH_OTSU)
        
        # Convert the processed image to PIL format for pytesseract
        pil_image = Image.fromarray(thresh)
        
        # Perform OCR
        text = pytesseract.image_to_string(pil_image)
        
        # Add metadata
        text = f"Image File: {os.path.basename(file_path)}\n\nExtracted Text:\n{text}"
        
        logger.info(f"Successfully extracted text from image: {file_path}")
        return text
    except Exception as e:
        logger.error(f"Error processing image {file_path}: {str(e)}")
        raise

def process_video(file_path):
    """
    Extract content from a video file by:
    1. Extracting frames at regular intervals
    2. Performing OCR on the extracted frames
    3. Extracting and transcribing the audio track
    
    Args:
        file_path (str): Path to the video file
        
    Returns:
        str: Extracted text content from video frames and audio
    """
    logger.info(f"Processing Video: {file_path}")
    try:
        # Load the video
        video = VideoFileClip(file_path)
        
        # Process video frames
        frame_interval = 5  # seconds
        frames = []
        for t in range(0, int(video.duration), frame_interval):
            frame = video.get_frame(t)
            # Convert from RGB to BGR (for OpenCV)
            frame = frame[:, :, ::-1].copy()
            frames.append(frame)
        
        logger.info(f"Extracted {len(frames)} frames from video: {file_path}")
        
        # Process frames with OCR
        frames_text = f"Video File: {os.path.basename(file_path)}\nDuration: {video.duration} seconds\n\nExtracted Text from Frames:\n"
        
        for i, frame in enumerate(frames):
            # Convert to grayscale
            gray = cv2.cvtColor(frame, cv2.COLOR_BGR2GRAY)
            
            # Apply thresholding
            _, thresh = cv2.threshold(gray, 150, 255, cv2.THRESH_BINARY | cv2.THRESH_OTSU)
            
            # Convert to PIL format for pytesseract
            pil_image = Image.fromarray(thresh)
            
            # Perform OCR
            frame_text = pytesseract.image_to_string(pil_image)
            
            if frame_text.strip():
                frames_text += f"\n--- Frame at {i * frame_interval} seconds ---\n{frame_text}"
        
        # Extract and process audio
        audio_text = "\n\nExtracted Text from Audio:\n"
        
        try:
            # Check if video has audio
            if video.audio is not None:
                logger.info(f"Extracting audio from video: {file_path}")
                
                # Save audio to a temporary file
                temp_audio_path = os.path.join(os.path.dirname(file_path), f"temp_audio_{os.path.basename(file_path)}.wav")
                video.audio.write_audiofile(temp_audio_path, logger=None)
                
                # Use SpeechRecognition for audio transcription
                try:
                    import speech_recognition as sr
                    recognizer = sr.Recognizer()
                    
                    # Process audio in chunks to handle longer files
                    from pydub import AudioSegment
                    from pydub.utils import make_chunks
                    
                    audio = AudioSegment.from_wav(temp_audio_path)
                    chunk_length_ms = 30000  # 30 seconds
                    chunks = make_chunks(audio, chunk_length_ms)
                    
                    logger.info(f"Processing {len(chunks)} audio chunks")
                    
                    for i, chunk in enumerate(chunks):
                        # Export chunk for processing
                        chunk_path = os.path.join(os.path.dirname(file_path), f"temp_chunk_{i}.wav")
                        chunk.export(chunk_path, format="wav")
                        
                        # Transcribe chunk
                        with sr.AudioFile(chunk_path) as source:
                            audio_data = recognizer.record(source)
                            try:
                                text = recognizer.recognize_google(audio_data)
                                if text:
                                    timestamp = i * (chunk_length_ms / 1000)
                                    audio_text += f"\n--- Audio at {timestamp:.1f} seconds ---\n{text}\n"
                            except sr.UnknownValueError:
                                pass
                            except sr.RequestError as e:
                                logger.warning(f"Google Speech Recognition service error: {e}")
                        
                        # Clean up chunk file
                        os.remove(chunk_path)
                
                except ImportError:
                    # Fallback to whisper if SpeechRecognition is not available
                    try:
                        import whisper
                        model = whisper.load_model("base")
                        result = model.transcribe(temp_audio_path)
                        audio_text += result["text"]
                    except ImportError:
                        logger.warning("Neither SpeechRecognition nor whisper is available for audio transcription")
                        audio_text += "Audio transcription unavailable - required libraries not installed."
                
                # Clean up temporary audio file
                os.remove(temp_audio_path)
            else:
                audio_text += "No audio track found in the video."
        except Exception as audio_error:
            logger.error(f"Error processing audio: {str(audio_error)}")
            audio_text += f"Error processing audio: {str(audio_error)}"
        
        # Combine frame and audio text
        all_text = frames_text + audio_text
        
        video.close()
        logger.info(f"Successfully extracted content from video: {file_path}")
        return all_text
    except Exception as e:
        logger.error(f"Error processing video {file_path}: {str(e)}")
        raise

def process_file(file_path):
    """
    Process a file based on its extension and extract text content.
    
    Args:
        file_path (str): Path to the file
        
    Returns:
        str: Extracted text content
    """
    # Get file extension
    _, file_extension = os.path.splitext(file_path)
    file_extension = file_extension.lower()
    
    # Process based on file type
    if file_extension == '.pdf':
        return process_pdf(file_path)
    elif file_extension == '.docx':
        return process_docx(file_path)
    elif file_extension == '.csv':
        return process_csv(file_path)
    elif file_extension in ['.jpg', '.jpeg', '.png', '.bmp', '.tiff']:
        return process_image(file_path)
    elif file_extension in ['.mp4', '.avi', '.mov', '.mkv', '.wmv']:
        return process_video(file_path)
    else:
        error_msg = f"Unsupported file format: {file_extension}"
        logger.error(error_msg)
        raise ValueError(error_msg)
